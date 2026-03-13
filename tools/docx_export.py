"""Export markdown article to a properly formatted DOCX file with SEO structure."""

from __future__ import annotations

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def markdown_to_docx(markdown_text: str, output_path: str, title: str = "") -> str:
    """
    Convert a markdown article to a DOCX file with proper heading hierarchy,
    lists, tables, bold/italic, and links.

    Returns the output file path.
    """
    doc = Document()

    # ── Style setup ──
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    for level in range(1, 4):
        heading_style = doc.styles[f"Heading {level}"]
        heading_style.font.name = "Calibri"
        heading_style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    doc.styles["Heading 1"].font.size = Pt(24)
    doc.styles["Heading 2"].font.size = Pt(18)
    doc.styles["Heading 3"].font.size = Pt(14)

    lines = markdown_text.split("\n")
    i = 0
    in_table = False
    table_rows = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # ── Skip empty lines ──
        if not stripped:
            i += 1
            continue

        # ── Headings ──
        if stripped.startswith("# ") and not stripped.startswith("## "):
            text = stripped[2:].strip()
            p = doc.add_heading(text, level=1)
            if not title:
                title = text
            i += 1
            continue

        if stripped.startswith("## "):
            text = stripped[3:].strip()
            doc.add_heading(text, level=2)
            i += 1
            continue

        if stripped.startswith("### "):
            text = stripped[4:].strip()
            doc.add_heading(text, level=3)
            i += 1
            continue

        # ── Table detection ──
        if "|" in stripped and stripped.startswith("|"):
            # Collect all table rows
            table_rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                row_text = lines[i].strip()
                # Skip separator rows (|---|---|)
                if re.match(r"^\|[\s\-:|]+\|$", row_text):
                    i += 1
                    continue
                cells = [c.strip() for c in row_text.strip("|").split("|")]
                table_rows.append(cells)
                i += 1

            if table_rows:
                _add_table(doc, table_rows)
            continue

        # ── Unordered lists ──
        if stripped.startswith("- ") or stripped.startswith("* "):
            bullet_items = []
            while i < len(lines) and (lines[i].strip().startswith("- ") or lines[i].strip().startswith("* ")):
                item_text = lines[i].strip()[2:].strip()
                bullet_items.append(item_text)
                i += 1
            for item in bullet_items:
                p = doc.add_paragraph(style="List Bullet")
                _add_rich_text(p, item)
            continue

        # ── Ordered lists ──
        if re.match(r"^\d+\.\s", stripped):
            list_items = []
            while i < len(lines) and re.match(r"^\d+\.\s", lines[i].strip()):
                item_text = re.sub(r"^\d+\.\s*", "", lines[i].strip())
                list_items.append(item_text)
                i += 1
            for item in list_items:
                p = doc.add_paragraph(style="List Number")
                _add_rich_text(p, item)
            continue

        # ── Horizontal rules ──
        if stripped in ("---", "***", "___"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            run = p.add_run("_" * 50)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            run.font.size = Pt(8)
            i += 1
            continue

        # ── Blockquotes ──
        if stripped.startswith("> "):
            quote_lines = []
            while i < len(lines) and lines[i].strip().startswith("> "):
                quote_lines.append(lines[i].strip()[2:])
                i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            _add_rich_text(p, " ".join(quote_lines))
            for run in p.runs:
                run.font.italic = True
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            continue

        # ── Regular paragraph ──
        p = doc.add_paragraph()
        _add_rich_text(p, stripped)
        i += 1

    # Save
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))
    return str(output)


def _add_table(doc: Document, rows: list[list[str]]):
    """Add a formatted table to the document."""
    if not rows:
        return

    num_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = "Light Grid Accent 1"

    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            if c_idx < num_cols:
                cell = table.rows[r_idx].cells[c_idx]
                cell.text = ""
                p = cell.paragraphs[0]
                _add_rich_text(p, cell_text)
                # Bold the header row
                if r_idx == 0:
                    for run in p.runs:
                        run.bold = True

    # Add spacing after table
    doc.add_paragraph()


def _add_rich_text(paragraph, text: str):
    """Parse markdown inline formatting and add runs to a paragraph.

    Handles: **bold**, *italic*, [link text](url), `code`
    """
    # Pattern to match inline markdown elements
    pattern = re.compile(
        r"(\*\*(.+?)\*\*)"       # bold
        r"|(\*(.+?)\*)"           # italic
        r"|(\[([^\]]+)\]\(([^)]+)\))"  # link
        r"|(`([^`]+)`)"           # inline code
    )

    pos = 0
    for match in pattern.finditer(text):
        # Add text before this match
        if match.start() > pos:
            paragraph.add_run(text[pos:match.start()])

        if match.group(2):  # bold
            run = paragraph.add_run(match.group(2))
            run.bold = True
        elif match.group(4):  # italic
            run = paragraph.add_run(match.group(4))
            run.italic = True
        elif match.group(6):  # link
            link_text = match.group(6)
            link_url = match.group(7)
            run = paragraph.add_run(link_text)
            run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
            run.underline = True
            # Store URL as a comment-style hint (DOCX doesn't natively support
            # hyperlinks through python-docx runs, but we can add them via the
            # underlying XML)
            _add_hyperlink(paragraph, link_text, link_url)
            # Remove the plain run we just added (hyperlink replaces it)
            paragraph._p.remove(run._r)
        elif match.group(9):  # inline code
            run = paragraph.add_run(match.group(9))
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x88, 0x00, 0x44)

        pos = match.end()

    # Add remaining text
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _add_hyperlink(paragraph, text: str, url: str):
    """Add a proper hyperlink to a paragraph using python-docx internals."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import hashlib

    # Generate a unique relationship ID
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    # Create the hyperlink element
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    # Create the run inside the hyperlink
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    # Style: blue, underlined
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)

    new_run.append(rPr)

    # Add the text
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    text_elem.set(qn("xml:space"), "preserve")
    new_run.append(text_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
